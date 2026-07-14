import os
import re
import html
import time
from pathlib import Path
from config import MINUTES_DIR

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <title>{title} - Minutes of Meeting</title>
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Lora:ital,wght@0,400;0,500;0,600;0,700;1,400&display=swap');
        
        :root {{
            --primary: #1E3A8A; /* Academic Navy Blue */
            --primary-light: #EFF6FF;
            --text-main: #0F172A;
            --text-muted: #334155;
            --bg-body: #F8FAFC;
            --bg-card: #FFFFFF;
            --border-color: #E2E8F0;
            --shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.05), 0 2px 4px -1px rgba(0, 0, 0, 0.025);
        }}

        body {{
            font-family: 'Lora', Georgia, serif;
            background-color: var(--bg-body);
            color: var(--text-main);
            margin: 0;
            padding: 40px 20px;
            line-height: 1.6;
        }}

        .container {{
            max-width: 800px;
            margin: 0 auto;
            background-color: var(--bg-card);
            border-radius: 8px;
            box-shadow: var(--shadow);
            overflow: hidden;
            border: 1px solid var(--border-color);
            border-top: 6px solid var(--primary);
        }}

        .header {{
            background: #FFFFFF;
            color: var(--text-main);
            padding: 40px 40px 30px 40px;
            position: relative;
            border-bottom: 2px solid var(--border-color);
        }}

        .badge {{
            background-color: var(--primary-light);
            color: var(--primary);
            padding: 6px 12px;
            border-radius: 4px;
            font-size: 0.75rem;
            font-weight: 700;
            text-transform: uppercase;
            letter-spacing: 0.05em;
            display: inline-block;
            margin-bottom: 16px;
        }}

        h1 {{
            margin: 0 0 16px 0;
            font-size: 1.8rem;
            font-weight: 700;
            letter-spacing: -0.025em;
            color: var(--text-main);
        }}

        .meta-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(180px, 1fr));
            gap: 16px;
            margin-top: 24px;
            font-size: 0.875rem;
            border-top: 1px solid var(--border-color);
            padding-top: 16px;
        }}

        .meta-item {{
            display: flex;
            flex-direction: column;
        }}

        .meta-label {{
            color: #64748B;
            font-size: 0.75rem;
            text-transform: uppercase;
            font-weight: 600;
            margin-bottom: 4px;
        }}

        .meta-value {{
            font-weight: 600;
            color: var(--text-main);
        }}

        .content {{
            padding: 40px;
        }}

        /* Style all levels of headers in content */
        h2 {{
            color: var(--primary);
            font-size: 1.3rem;
            font-weight: 600;
            margin-top: 32px;
            margin-bottom: 16px;
            border-bottom: 2px solid var(--primary-light);
            padding-bottom: 8px;
            display: flex;
            align-items: center;
            gap: 8px;
        }}

        h2::before {{
            content: "";
            display: inline-block;
            width: 4px;
            height: 18px;
            background-color: var(--primary);
            border-radius: 2px;
        }}

        h3 {{
            color: var(--primary);
            font-size: 1.15rem;
            font-weight: 600;
            margin-top: 28px;
            margin-bottom: 12px;
            display: flex;
            align-items: center;
            gap: 8px;
        }}

        h3::before {{
            content: "";
            display: inline-block;
            width: 4px;
            height: 15px;
            background-color: var(--primary);
            border-radius: 2px;
        }}

        p {{
            margin: 0 0 16px 0;
            color: var(--text-muted);
        }}

        ul, ol {{
            margin: 0 0 24px 0;
            padding-left: 20px;
            color: var(--text-muted);
        }}

        li {{
            margin-bottom: 8px;
        }}

        /* Checkbox Lists */
        .task-list {{
            list-style-type: none;
            padding-left: 0;
        }}

        .task-item {{
            display: flex;
            align-items: flex-start;
            gap: 12px;
            margin-bottom: 10px;
            padding: 10px 14px;
            border-radius: 8px;
            background-color: #F9FAFB;
            border: 1px solid var(--border-color);
        }}

        .checkbox {{
            display: inline-flex;
            align-items: center;
            justify-content: center;
            width: 18px;
            height: 18px;
            border: 2px solid #D1D5DB;
            border-radius: 4px;
            background-color: #FFFFFF;
            flex-shrink: 0;
            margin-top: 3px;
        }}

        .checkbox.checked {{
            background-color: var(--primary);
            border-color: var(--primary);
            color: #FFFFFF;
            font-size: 11px;
            font-weight: bold;
        }}

        .task-item.checked {{
            text-decoration: line-through;
            color: #9CA3AF;
            background-color: #F3F4F6;
        }}

        .institution {{
            font-size: 1.15rem;
            font-weight: 800;
            letter-spacing: 0.03em;
            text-transform: uppercase;
            color: var(--primary);
            margin-bottom: 4px;
        }}

        .center-name {{
            font-size: 0.85rem;
            font-weight: 700;
            letter-spacing: 0.06em;
            text-transform: uppercase;
            color: #64748B;
            margin-bottom: 24px;
            border-bottom: 1px solid var(--border-color);
            padding-bottom: 8px;
        }}

        footer {{
            text-align: center;
            padding: 24px;
            font-size: 0.75rem;
            color: #9CA3AF;
            border-top: 1px solid var(--border-color);
            background-color: #F9FAFB;
            line-height: 1.5;
        }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <div class="institution">K S RANGASAMY COLLEGE OF TECHNOLOGY</div>
            <div class="center-name">TRAIT CENTER</div>
            <span class="badge">Minutes of Meeting</span>
            <h1>{title}</h1>
            <div class="meta-grid">
                <div class="meta-item">
                    <span class="meta-label">Meeting ID</span>
                    <span class="meta-value">#{meeting_id}</span>
                </div>
                <div class="meta-item">
                    <span class="meta-label">Date & Time</span>
                    <span class="meta-value">{date_time}</span>
                </div>
            </div>
        </div>
        
        <div class="content">
            {html_content}
        </div>

        <footer>
            Generated automatically by Meet Meeting Assistant on {gen_time}<br>
            <strong>K S Rangasamy College of Technology — TRAIT Center</strong>
        </footer>
    </div>
</body>
</html>
"""


def markdown_to_html(md_text: str) -> str:
    """A lightweight, robust parser to convert standard Markdown to clean HTML elements."""
    escaped_text = html.escape(md_text)
    lines = escaped_text.splitlines()
    html_lines = []
    
    in_list = False
    in_ordered_list = False
    
    for line in lines:
        stripped = line.strip()
        if not stripped:
            if in_list:
                html_lines.append("</ul>")
                in_list = False
            if in_ordered_list:
                html_lines.append("</ol>")
                in_ordered_list = False
            continue
            
        # Parse headings (### or ## or #)
        heading_match = re.match(r'^(#{1,6})\s+(.*)', stripped)
        if heading_match:
            if in_list:
                html_lines.append("</ul>")
                in_list = False
            if in_ordered_list:
                html_lines.append("</ol>")
                in_ordered_list = False
            level = len(heading_match.group(1))
            # Increase heading levels by 1 for cleaner document structure (e.g. h2/h3 instead of h1/h2)
            html_level = min(level + 1, 6)
            content = heading_match.group(2)
            content = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', content)
            html_lines.append(f"<h{html_level}>{content}</h{html_level}>")
            continue
            
        # Parse checkboxes (e.g. - [ ] task)
        chk_match = re.match(r'^[-*+]\s+\[\s*\]\s+(.*)', stripped)
        if chk_match:
            if not in_list:
                if in_ordered_list:
                    html_lines.append("</ol>")
                    in_ordered_list = False
                html_lines.append("<ul class='task-list'>")
                in_list = True
            content = chk_match.group(1)
            content = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', content)
            html_lines.append(f"<li class='task-item'><span class='checkbox'></span><span>{content}</span></li>")
            continue
            
        # Parse checked checkboxes (e.g. - [x] task)
        chk_x_match = re.match(r'^[-*+]\s+\[[xX]\]\s+(.*)', stripped)
        if chk_x_match:
            if not in_list:
                if in_ordered_list:
                    html_lines.append("</ol>")
                    in_ordered_list = False
                html_lines.append("<ul class='task-list'>")
                in_list = True
            content = chk_x_match.group(1)
            content = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', content)
            html_lines.append(f"<li class='task-item checked'><span class='checkbox checked'>✓</span><span>{content}</span></li>")
            continue

        # Parse bullet lists
        bullet_match = re.match(r'^[-*+]\s+(.*)', stripped)
        if bullet_match:
            if not in_list:
                if in_ordered_list:
                    html_lines.append("</ol>")
                    in_ordered_list = False
                html_lines.append("<ul>")
                in_list = True
            content = bullet_match.group(1)
            content = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', content)
            html_lines.append(f"<li>{content}</li>")
            continue
            
        # Parse ordered lists
        ol_match = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if ol_match:
            if not in_ordered_list:
                if in_list:
                    html_lines.append("</ul>")
                    in_list = False
                html_lines.append("<ol>")
                in_ordered_list = True
            content = ol_match.group(2)
            content = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', content)
            html_lines.append(f"<li>{content}</li>")
            continue
            
        # Normal lines (paragraphs)
        if in_list:
            html_lines.append("</ul>")
            in_list = False
        if in_ordered_list:
            html_lines.append("</ol>")
            in_ordered_list = False
            
        content = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', stripped)
        html_lines.append(f"<p>{content}</p>")
        
    if in_list:
        html_lines.append("</ul>")
    if in_ordered_list:
        html_lines.append("</ol>")
        
    return "\n".join(html_lines)


def _generate_docx(docx_path: str, title: str, meeting_id: int, date_time_str: str, gen_time_str: str, mom_content: str):
    doc = Document()
    
    # Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Helper to set font of a run
    def set_font(run, font_name='Lora', size_pt=11, color_rgb=(15, 23, 42), bold=False, italic=False):
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run.font.color.rgb = RGBColor(*color_rgb)
        run.bold = bold
        run.italic = italic
        
    # Helper to add bottom border
    def add_bottom_border(paragraph, color_hex="CBD5E1", size_val="12"):
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), size_val)
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), color_hex)
        pBdr.append(bottom)
        pPr.append(pBdr)

    # 1. LETTERHEAD
    p_inst = doc.add_paragraph()
    p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_inst.paragraph_format.space_after = Pt(2)
    run_inst = p_inst.add_run("K S RANGASAMY COLLEGE OF TECHNOLOGY")
    set_font(run_inst, font_name='Lora', size_pt=14, color_rgb=(30, 58, 138), bold=True)
    
    p_center = doc.add_paragraph()
    p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_center.paragraph_format.space_after = Pt(12)
    run_center = p_center.add_run("TRAIT CENTER")
    set_font(run_center, font_name='Lora', size_pt=10, color_rgb=(100, 116, 139), bold=True)
    add_bottom_border(p_center, color_hex="E2E8F0", size_val="12")
    
    # 2. DOCUMENT TITLE
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(18)
    run_title = p_title.add_run("MINUTES OF MEETING REPORT")
    set_font(run_title, font_name='Lora', size_pt=12, color_rgb=(15, 23, 42), bold=True)
    
    # 3. METADATA TABLE (2x2)
    table = doc.add_table(rows=2, cols=2)
    table.autofit = False
    
    # Set column widths
    table.columns[0].width = Inches(3.25)
    table.columns[1].width = Inches(3.25)
    
    # Row 0
    cell_title = table.cell(0, 0)
    p_cell_t = cell_title.paragraphs[0]
    p_cell_t.paragraph_format.space_after = Pt(4)
    run_t_lbl = p_cell_t.add_run("Meeting Title: ")
    set_font(run_t_lbl, font_name='Lora', size_pt=9.5, color_rgb=(100, 116, 139), bold=True)
    run_t_val = p_cell_t.add_run(title)
    set_font(run_t_val, font_name='Lora', size_pt=9.5, color_rgb=(15, 23, 42), bold=False)
    
    cell_date = table.cell(0, 1)
    p_cell_d = cell_date.paragraphs[0]
    p_cell_d.paragraph_format.space_after = Pt(4)
    run_d_lbl = p_cell_d.add_run("Date & Time: ")
    set_font(run_d_lbl, font_name='Lora', size_pt=9.5, color_rgb=(100, 116, 139), bold=True)
    run_d_val = p_cell_d.add_run(date_time_str)
    set_font(run_d_val, font_name='Lora', size_pt=9.5, color_rgb=(15, 23, 42), bold=False)
    
    # Row 1
    cell_id = table.cell(1, 0)
    p_cell_id = cell_id.paragraphs[0]
    p_cell_id.paragraph_format.space_after = Pt(4)
    run_id_lbl = p_cell_id.add_run("Meeting ID: ")
    set_font(run_id_lbl, font_name='Lora', size_pt=9.5, color_rgb=(100, 116, 139), bold=True)
    run_id_val = p_cell_id.add_run(f"#{meeting_id}")
    set_font(run_id_val, font_name='Lora', size_pt=9.5, color_rgb=(15, 23, 42), bold=False)
    
    cell_gen = table.cell(1, 1)
    p_cell_gen = cell_gen.paragraphs[0]
    p_cell_gen.paragraph_format.space_after = Pt(4)
    run_gen_lbl = p_cell_gen.add_run("Generated On: ")
    set_font(run_gen_lbl, font_name='Lora', size_pt=9.5, color_rgb=(100, 116, 139), bold=True)
    run_gen_val = p_cell_gen.add_run(gen_time_str)
    set_font(run_gen_val, font_name='Lora', size_pt=9.5, color_rgb=(15, 23, 42), bold=False)
    
    # Add spacing after table
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_before = Pt(12)
    p_spacer.paragraph_format.space_after = Pt(6)
    add_bottom_border(p_spacer, color_hex="E2E8F0", size_val="12")
    
    # 4. CONTENT SECTIONS
    lines = mom_content.splitlines()
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
            
        # Match headings (e.g. ### 1. Summary)
        heading_match = re.match(r'^(#{1,6})\s+(.*)', stripped)
        if heading_match:
            header_text = heading_match.group(2)
            header_text = re.sub(r'\*\*(.*?)\*\*', r'\1', header_text)
            
            p_head = doc.add_paragraph()
            p_head.paragraph_format.space_before = Pt(18)
            p_head.paragraph_format.space_after = Pt(6)
            p_head.paragraph_format.keep_with_next = True
            
            run_head = p_head.add_run(header_text)
            set_font(run_head, font_name='Lora', size_pt=12, color_rgb=(30, 58, 138), bold=True)
            add_bottom_border(p_head, color_hex="EFF6FF", size_val="6")
            continue
            
        # Parse items (bullets, checklists)
        p_item = doc.add_paragraph()
        p_item.paragraph_format.space_after = Pt(4)
        
        # Checked action items
        if stripped.startswith("- [x]") or stripped.startswith("- [X]"):
            content = stripped[5:].strip()
            content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)
            run_check = p_item.add_run("☑  ")
            set_font(run_check, font_name='Lora', size_pt=10.5, color_rgb=(30, 58, 138), bold=True)
            run_text = p_item.add_run(content)
            set_font(run_text, font_name='Lora', size_pt=10.5, color_rgb=(100, 116, 139))
            run_text.font.strike = True
            p_item.paragraph_format.left_indent = Inches(0.25)
            
        # Unchecked action items
        elif stripped.startswith("- [ ]"):
            content = stripped[5:].strip()
            content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)
            run_check = p_item.add_run("☐  ")
            set_font(run_check, font_name='Lora', size_pt=10.5, color_rgb=(30, 58, 138), bold=True)
            run_text = p_item.add_run(content)
            set_font(run_text, font_name='Lora', size_pt=10.5, color_rgb=(15, 23, 42))
            p_item.paragraph_format.left_indent = Inches(0.25)
            
        # Bullets
        elif stripped.startswith("- "):
            content = stripped[2:].strip()
            content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)
            run_bullet = p_item.add_run("•  ")
            set_font(run_bullet, font_name='Lora', size_pt=10.5, color_rgb=(30, 58, 138), bold=True)
            run_text = p_item.add_run(content)
            set_font(run_text, font_name='Lora', size_pt=10.5, color_rgb=(51, 65, 85))
            p_item.paragraph_format.left_indent = Inches(0.25)
            
        # Numbered list
        elif re.match(r'^\d+\.\s+', stripped):
            num_match = re.match(r'^(\d+\.)\s+(.*)', stripped)
            num_prefix = num_match.group(1)
            content = num_match.group(2)
            content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)
            run_num = p_item.add_run(f"{num_prefix}  ")
            set_font(run_num, font_name='Lora', size_pt=10.5, color_rgb=(30, 58, 138), bold=True)
            run_text = p_item.add_run(content)
            set_font(run_text, font_name='Lora', size_pt=10.5, color_rgb=(51, 65, 85))
            p_item.paragraph_format.left_indent = Inches(0.25)
            
        # Normal paragraph text
        else:
            content = re.sub(r'\*\*(.*?)\*\*', r'\1', stripped)
            run_text = p_item.add_run(content)
            set_font(run_text, font_name='Lora', size_pt=10.5, color_rgb=(51, 65, 85))
            p_item.paragraph_format.space_after = Pt(8)
            
    # 5. FOOTER SECTION
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_foot.paragraph_format.space_before = Pt(36)
    add_bottom_border(p_foot, color_hex="E2E8F0", size_val="6")
    
    p_foot_text = doc.add_paragraph()
    p_foot_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_foot_text.paragraph_format.space_before = Pt(6)
    
    run_f1 = p_foot_text.add_run(f"Generated automatically by Meet Meeting Assistant on {gen_time_str}\n")
    set_font(run_f1, font_name='Lora', size_pt=8.5, color_rgb=(148, 163, 184))
    
    run_f2 = p_foot_text.add_run("K S Rangasamy College of Technology — TRAIT Center")
    set_font(run_f2, font_name='Lora', size_pt=8.5, color_rgb=(100, 116, 139), bold=True)
    
    doc.save(docx_path)


def generate_minutes_documents(meeting_id: int, title: str, start_time: float, mom_content: str) -> tuple[str, str, str]:
    """Generates .txt (memo format), styled .html, and styled Microsoft Word .docx Minutes of Meeting documents.
    
    Returns:
        tuple[str, str, str]: Absolute paths to (text_file, html_file, docx_file)
    """
    os.makedirs(MINUTES_DIR, exist_ok=True)
    
    # Safe filenames
    safe_title = re.sub(r'[^a-zA-Z0-9_\-]', '_', title)
    txt_filename = f"Minutes_of_Meeting_ID_{meeting_id}_{safe_title}.txt"
    html_filename = f"Minutes_of_Meeting_ID_{meeting_id}_{safe_title}.html"
    docx_filename = f"Minutes_of_Meeting_ID_{meeting_id}_{safe_title}.docx"
    
    txt_path = os.path.join(MINUTES_DIR, txt_filename)
    html_path = os.path.join(MINUTES_DIR, html_filename)
    docx_path = os.path.join(MINUTES_DIR, docx_filename)
    
    # Format metadata
    date_time_str = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(start_time))
    gen_time_str = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime())
    
    # 1. Generate beautifully formatted plain text layout
    border = "-" * 80
    
    inst_line = "K S RANGASAMY COLLEGE OF TECHNOLOGY"
    inst_padding = (80 - len(inst_line)) // 2
    inst_formatted = " " * max(0, inst_padding) + inst_line
    
    center_line = "TRAIT CENTER"
    center_padding = (80 - len(center_line)) // 2
    center_formatted = " " * max(0, center_padding) + center_line
    
    title_line = "MINUTES OF MEETING REPORT"
    title_padding = (80 - len(title_line)) // 2
    title_formatted = " " * max(0, title_padding) + title_line
    
    txt_content = f"{inst_formatted}\n"
    txt_content += f"{center_formatted}\n"
    txt_content += f"{border}\n"
    txt_content += f"{title_formatted}\n"
    txt_content += f"{border}\n\n"
    
    title_lbl = f"Meeting Title : {title}"
    id_lbl    = f"Meeting ID    : #{meeting_id}"
    date_lbl  = f"Date & Time   : {date_time_str}"
    gen_lbl   = f"Generated On  : {gen_time_str}"
    
    txt_content += f"  {title_lbl:<44} {date_lbl}\n"
    txt_content += f"  {id_lbl:<44} {gen_lbl}\n"
    txt_content += f"{border}\n\n"
    
    # Parse headers and format lists/todos cleanly
    formatted_mom = ""
    lines = mom_content.splitlines()
    for line in lines:
        stripped = line.strip()
        if not stripped:
            formatted_mom += "\n"
            continue
            
        # Match headings (e.g. ### 1. Summary)
        heading_match = re.match(r'^(#{1,6})\s+(.*)', stripped)
        if heading_match:
            header_text = heading_match.group(2)
            # Remove bold markdown from headers
            header_text = re.sub(r'\*\*(.*?)\*\*', r'\1', header_text)
            formatted_mom += f"\n[{header_text.upper()}]\n{border}\n"
        else:
            # Clean up bold markers (**text**) and other md features for plain text representation
            clean_line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            # Format list bullets nicely if they are simple markdown
            if clean_line.strip().startswith("- [ ]"):
                clean_line = "  [ ] " + clean_line.strip()[5:].strip()
            elif clean_line.strip().startswith("- [x]") or clean_line.strip().startswith("- [X]"):
                clean_line = "  [x] " + clean_line.strip()[5:].strip()
            elif clean_line.strip().startswith("- "):
                clean_line = "  • " + clean_line.strip()[2:].strip()
            elif re.match(r'^\d+\.\s+', clean_line.strip()):
                clean_line = "  " + clean_line.strip()
                
            formatted_mom += clean_line + "\n"
            
    txt_content += formatted_mom
    txt_content += f"\n{border}\n"
    txt_content += f"{' ' * 23}Generated by Meet Meeting Assistant\n"
    txt_content += f"{' ' * 13}K S Rangasamy College of Technology - TRAIT Center\n"
    txt_content += f"{border}\n"
    
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(txt_content)
        
    # 2. Generate HTML File
    parsed_html_content = markdown_to_html(mom_content)
    html_full = HTML_TEMPLATE.format(
        title=html.escape(title),
        meeting_id=meeting_id,
        date_time=date_time_str,
        html_content=parsed_html_content,
        gen_time=gen_time_str
    )
    
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html_full)
        
    # 3. Generate Microsoft Word Document (.docx)
    try:
        _generate_docx(docx_path, title, meeting_id, date_time_str, gen_time_str, mom_content)
    except Exception as e:
        print(f"[DocGenerator] Error generating Word docx: {e}")
        
    print(f"[DocGenerator] Generated documents:\n  - {txt_path}\n  - {html_path}\n  - {docx_path}")
    return txt_path, html_path, docx_path
